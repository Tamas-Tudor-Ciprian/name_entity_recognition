from ner import *

from docx import Document




def run():
    input_doc = Document(r"C:\Users\Tudor\Desktop\aaa.docx")

    name_list = name_filter(input_doc)

    print("Filtering done!")

    output_doc = Document()
    for index ,name in enumerate(name_list):
        output_doc.add_paragraph(f"{index+1}.{name}")

    print("Output file created!")
    output_doc.save(r"C:\Users\Tudor\Desktop\output.docx")

    print("Export done!")


if __name__ == "__main__":
    run()